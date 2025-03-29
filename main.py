import csv
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

def set_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_line(paragraph, length, thickness):
    run = paragraph.add_run('─' * length)
    run.font.size = Pt(thickness)
    run.font.color.rgb = RGBColor(0, 0, 0)

def create_art_labels(data_path, output_path):
    # Load the CSV file
    data = pd.read_csv(data_path)
    
    # Filter the rows where the 'Label' column is 'Y'
    data = data[data['Label'] == 'Y']

    # Create a new Word document
    doc = Document()
    
    # Set the margins to match Avery 5371 template
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set up the grid: 2 columns, 5 rows per page (10 labels per page)
    labels_per_page = 10
    labels_per_row = 2
    
    # Loop through the data and create labels
    for i, row in data.iterrows():
        if i % labels_per_page == 0 and i > 0:
            doc.add_page_break()
        
        # Add the table for the grid layout
        if i % labels_per_row == 0:
            table = doc.add_table(rows=1, cols=labels_per_row)
            table.autofit = False
            table.allow_autofit = False
            for col in table.columns:
                col.width = Inches(3.75)
        
        # Populate the label
        artist, title, medium, size, price = row['Artist'], row['Title'], row['Medium'], row['Size'], row['Price']
        cell = table.cell(0, i % labels_per_row)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Add a new paragraph for each label
        p = cell.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the line above the artist
        add_line(p, 125, 2)  
        p.add_run('\n')

        # Add the artist
        artist_run = p.add_run(f'{artist}\n')
        set_font(artist_run, 'Didot', 14, True)

        # Add the line below the artist
        add_line(p, 80, 4)  
        p.add_run('\n')

        # Add the title
        title_run = p.add_run(f'{title}\n')
        set_font(title_run, 'Avenir Next Regular', 15, True)

        # Add the medium
        if medium:
            medium_run = p.add_run(f'{medium}\n')
            set_font(medium_run, 'Hoefler Text', 10, True)

        # Add the size
        if size:
            size_run = p.add_run(f'{size}\n')
            set_font(size_run, 'Hoefler Text', 10, True)

        # Add the price
        if price:
            price_run = p.add_run(f'{price}\n')
            set_font(price_run, 'Hoefler Text', 14, True)
        
        # Center-align the text in the cell
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
    # Save the document
    doc.save(output_path)

# Paths to the CSV file and the output Word document
csv_path = "Assets - Product (Sellable Items).csv"
output_doc_path = "Generated_Art_Labels.docx"

# Generate the labels
create_art_labels(csv_path, output_doc_path)

print(f"Labels have been successfully saved to {output_doc_path}")